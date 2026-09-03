import json
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "website-content-review.docx"


def get_data():
    command = (
        "const fs=require('fs'),vm=require('vm');"
        "const sandbox={};"
        "vm.runInNewContext(fs.readFileSync('assets/js/data.js','utf8')"
        "+'\\nthis.exportedData=CV_DATA;',sandbox);"
        "process.stdout.write(JSON.stringify(sandbox.exportedData));"
    )
    result = subprocess.run(["node", "-e", command], cwd=ROOT, capture_output=True, text=True, check=True)
    return json.loads(result.stdout)


def set_rtl(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(6)
    properties = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    properties.append(bidi)


def shade(cell, color):
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    properties.append(shading)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_rtl(paragraph)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_text(document, text, bold=False, size=10, color=None, indent=0):
    paragraph = document.add_paragraph()
    set_rtl(paragraph)
    paragraph.paragraph_format.right_indent = Cm(indent)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return paragraph


def add_heading(document, text, level=1):
    heading = document.add_heading()
    set_rtl(heading)
    heading.style = f"Heading {min(level, 3)}"
    run = heading.add_run(text)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(51, 75, 184) if level == 1 else RGBColor(24, 36, 58)
    return heading


def add_review_box(document):
    table = document.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.style = "Table Grid"
    labels = ["נוסח חדש / תיקון", "הערת תוכן או עובדה", "הערה ויזואלית / UX"]
    for row, label in zip(table.rows, labels):
        set_cell_text(row.cells[0], label, bold=True, color=(51, 75, 184))
        shade(row.cells[0], "EAF1F8")
        set_cell_text(row.cells[1], "", color=(100, 116, 139))
        row.cells[1].height = Cm(1.05)
    document.add_paragraph()


def add_links(document, links):
    if not links:
        return
    add_text(document, "רפרנסים וקישורים באתר", bold=True, size=10, color=(51, 75, 184))
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.style = "Table Grid"
    for cell, title in zip(table.rows[0].cells, ["מקור", "כותרת / תיאור", "קישור"]):
        set_cell_text(cell, title, bold=True, color=(255, 255, 255))
        shade(cell, "334BB8")
    for link in links:
        cells = table.add_row().cells
        set_cell_text(cells[0], link.get("source", ""))
        set_cell_text(cells[1], f"{link.get('title', '')}\n{link.get('desc', link.get('subtitle', ''))}")
        set_cell_text(cells[2], link.get("url", ""), color=(23, 138, 192))
    document.add_paragraph()


def add_experience(document, item):
    add_heading(document, f"{item['period']} | {item['role']}", 2)
    add_text(document, f"{item['company']} | {item['companyType']}", bold=True, color=(21, 147, 106))
    add_text(document, f"תגית באתר: {item.get('badge', '')}", size=9, color=(100, 116, 139))
    if item.get("directReport"):
        add_text(document, item["directReport"], size=9, color=(100, 116, 139))
    add_text(document, "תמצית מנהלים באתר", bold=True, size=10)
    add_text(document, item["shortSummary"], size=10)
    add_text(document, "פירוט מלא (נפתח בלחיצה באתר)", bold=True, size=10)
    for bullet in item["bullets"]:
        add_text(document, f"• {bullet}", size=10, indent=0.35)
    add_text(document, f"תגיות באתר: {', '.join(item.get('tags', []))}", size=9, color=(100, 116, 139))
    add_links(document, item.get("links", []))
    add_review_box(document)


def main():
    data = get_data()
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    title = document.add_heading()
    set_rtl(title)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("חוברת ביקורת ועדכון תוכן האתר")
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(51, 75, 184)
    subtitle = add_text(document, "אמיר רוזן | קורות חיים ופורטפוליו מנהלים", size=14, color=(24, 36, 58))
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(document, "מטרת המסמך: לעבור בנוחות על כל הטקסטים, הכותרות, הרפרנסים והחלטות ה-UX באתר. אפשר למחוק, לערוך ולהוסיף הערות ישירות ב-Word. לאחר קבלת הקובץ המתוקן אעדכן את האתר בהתאם.", size=11)

    add_heading(document, "איך עובדים עם המסמך", 1)
    for instruction in [
        "לשינוי טקסט: ערכו ישירות את הטקסט המופיע במסמך, או כתבו נוסח חלופי בשדה 'נוסח חדש / תיקון'.",
        "להערת תוכן: כתבו עובדה שצריך לאמת, השלמה, קיצור או ניסוח רצוי בשדה 'הערת תוכן או עובדה'.",
        "להערת עיצוב: כתבו מה תרצו לראות או להרגיש בשדה 'הערה ויזואלית / UX'.",
        "אין צורך להתייחס לקוד או לצלם מסכים. אפשר להחזיר את הקובץ המתוקן כמות שהוא.",
    ]:
        add_text(document, f"• {instruction}")

    add_heading(document, "1. פתיח, זהות ופרטי קשר", 1)
    personal = data["personal"]
    rows = [
        ("שם", personal["name"]),
        ("שם באנגלית", personal["englishName"]),
        ("כותרת ראשית", personal["title"]),
        ("תג הישג בפתיח", "מקום 1 בישראל באימוץ AI (מדד Evident AI 2026) | בנק לאומי"),
        ("תקציר פתיח באתר", "מנהל בכיר בעל מעל 20 שנות ניסיון רב-תחומי, המשלב ראייה אסטרטגית עסקית רחבה עם מומחיות טכנולוגית מעמיקה (כולל Hands-on). הובלת מרכז ה-AI של בנק לאומי בכפיפות ישירה למנכ״ל, ניהול אגפים ויחידות P&L בהיקפי עשרות מיליוני שקלים, ניהול קרן הון סיכון (Irani Ventures), ודירקטוריונים."),
        ("טלפון", personal["phone"]),
        ("דוא״ל", personal["email"]),
        ("LinkedIn", personal["linkedin"]),
        ("קישור לפרק אישי", "קראו גם על האדם שמאחורי המסע | משפחה, ערכים והדרך האישית"),
    ]
    for label, value in rows:
        add_text(document, label, bold=True, size=10, color=(51, 75, 184))
        add_text(document, value, size=10)
        add_review_box(document)

    add_heading(document, "2. מדדי אימפקט", 1)
    for stat in personal["stats"]:
        add_text(document, f"{stat['value']} | {stat['label']}", bold=True, size=11)
        add_text(document, stat["sublabel"], size=10)
        add_review_box(document)

    add_heading(document, "3. יכולות ליבה", 1)
    for competency in data["competencies"]:
        add_text(document, f"{competency['name']} | {competency['level']}%", bold=True, size=11)
        add_text(document, competency["desc"], size=10)
        add_review_box(document)

    add_heading(document, "4. המסע המקצועי", 1)
    add_text(document, "באתר כל תפקיד מוצג בציר זמן. הכותרת והתמצית מוצגות מיד; הפירוט והרפרנסים נפתחים בלחיצה.", size=10, color=(100, 116, 139))
    for item in data["experience"]:
        add_experience(document, item)

    add_heading(document, "5. סיקור תקשורתי, מדיה ורפרנסים", 1)
    add_text(document, "רשימת כל הפריטים המופיעים במרכז המדיה של האתר. בדקו ניסוח, שיוך לתפקיד ורלוונטיות של כל קישור.", size=10)
    add_links(document, data["mediaHub"])
    add_review_box(document)

    add_heading(document, "6. השכלה, אקדמיה ומנטורינג", 1)
    for item in data["educationAndMentorship"]["education"]:
        add_text(document, f"{item['period']} | {item['degree']}", bold=True, size=11)
        add_text(document, f"{item['institution']} - {item['details']}", size=10)
        add_review_box(document)
    for item in data["educationAndMentorship"]["academicAndMentoring"]:
        add_text(document, f"{item['period']} | {item['role']}", bold=True, size=11)
        add_text(document, f"{item['institution']} - {item['details']}", size=10)
        add_links(document, item.get("links", []))
        add_review_box(document)

    add_heading(document, "7. הפרק האישי והמשפחה", 1)
    add_text(document, "כותרת באתר: העוגן שלי - משפחה, ערכים ותשוקה", bold=True, size=11)
    for paragraph in personal["personalStory"]["paragraphs"]:
        add_text(document, paragraph, size=10)
        add_review_box(document)
    add_text(document, "נקודות לחיצות על התמונה", bold=True, size=11, color=(51, 75, 184))
    family = [
        ("שלי", "מטפלת זוגית ופסיכותרפיסטית. הכרנו בצבא, ומאז אנחנו יחד באושר."),
        ("יונתן", "הבן המרכזי שלנו, בן 16. לומד בכיתת נחשון ומבלה את רוב זמנו בצופים."),
        ("אני", "מחייך, ומודה בכל יום על המשפחה שהיא העוגן והכוח שמאחורי הדרך."),
        ("ליבי", "ליבי הבן, בן 10 - הצעיר שלנו, עם שם שאנשים נוטים בטעות לייחס לבת."),
        ("גל", "בן 18, מתגייס לחיל האוויר בדצמבר. המינימי שלי."),
    ]
    for name, description in family:
        add_text(document, name, bold=True, size=11)
        add_text(document, description, size=10)
        add_review_box(document)
    add_text(document, "ערכים אישיים באתר", bold=True, size=11, color=(51, 75, 184))
    for value in personal["personalStory"]["values"]:
        add_text(document, value["title"], bold=True, size=10)
        add_text(document, value["desc"], size=10)
        add_review_box(document)

    add_heading(document, "8. רכיבי ממשק וחוויה ויזואלית", 1)
    components = [
        "סגנון כללי: עיצוב בהיר, תכלת ואינדיגו, ציר זמן מקצועי, מותאם מובייל תחילה.",
        "פתיח: שם, תפקיד, תקציר, פרטי קשר וקישור לפרק האישי.",
        "ציר זמן: תחנות לחיצות, סינון, חיפוש, פירוט מתקפל באמצעות + / - ורפרנסים לכל תפקיד.",
        "מדיה: כרטיסי כתבות, פוסטים וראיונות עם סינון.",
        "תמונה משפחתית: נקודות לחיצות עם +, בועת טקסט שקופה בראש התמונה וכפתור סגירה X.",
        "ניווט מובייל תחתון: ראשי, ניסיון, מדיה, אישי, קשר.",
        "אנימציות: חשיפה מדורגת בגלילה, קפיצה בין תחנות בציר הזמן, תנועת רקע עדינה ואנימציות מגע קטנות.",
    ]
    for component in components:
        add_text(document, component, size=10)
        add_review_box(document)

    add_heading(document, "9. הערות כלליות", 1)
    for _ in range(8):
        add_review_box(document)

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()